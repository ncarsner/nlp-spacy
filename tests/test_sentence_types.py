import sys
import os
from unittest.mock import Mock, patch, MagicMock
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from utils.sentence_types import read_text_from_file


class TestReadTextFromFile:
    """Test suite for read_text_from_file() in sentence_types.py"""

    def test_read_txt_file(self, tmp_path):
        """Test reading a plain text file with UTF-8 encoding"""
        txt_file = tmp_path / "sample.txt"
        txt_file.write_text("Hello world. How are you?", encoding='utf-8')
        result = read_text_from_file(str(txt_file))
        assert result == "Hello world. How are you?"

    def test_read_txt_file_utf8_content(self, tmp_path):
        """Test reading a plain text file containing non-ASCII UTF-8 characters"""
        txt_file = tmp_path / "utf8.txt"
        txt_file.write_text("Héllo wörld!", encoding='utf-8')
        result = read_text_from_file(str(txt_file))
        assert result == "Héllo wörld!"

    def test_file_not_found_exits(self, tmp_path):
        """Test that a missing file causes SystemExit"""
        with pytest.raises(SystemExit):
            read_text_from_file(str(tmp_path / "nonexistent.txt"))

    @patch('utils.sentence_types.Path')
    def test_read_docx_file(self, mock_path_cls):
        """Test reading a .docx file extracts paragraph text"""
        mock_path = MagicMock()
        mock_path.exists.return_value = True
        mock_path.suffix = '.docx'
        mock_path.__str__ = lambda self: 'document.docx'
        mock_path_cls.return_value = mock_path

        mock_para1 = Mock()
        mock_para1.text = "Hello world."
        mock_para2 = Mock()
        mock_para2.text = "How are you?"

        mock_doc = Mock()
        mock_doc.paragraphs = [mock_para1, mock_para2]

        with patch('utils.sentence_types.docx.Document', return_value=mock_doc):
            result = read_text_from_file('document.docx')

        assert result == "Hello world.\nHow are you?"

    @patch('utils.sentence_types.Path')
    def test_read_docx_error_exits(self, mock_path_cls):
        """Test that a docx reading error causes SystemExit"""
        mock_path = MagicMock()
        mock_path.exists.return_value = True
        mock_path.suffix = '.docx'
        mock_path.__str__ = lambda self: 'bad.docx'
        mock_path_cls.return_value = mock_path

        with patch('utils.sentence_types.docx.Document', side_effect=Exception("corrupt file")):
            with pytest.raises(SystemExit):
                read_text_from_file('bad.docx')

    def test_read_docx_real_file(self, tmp_path):
        """Test reading an actual .docx file created with python-docx"""
        import docx as docx_lib
        docx_file = tmp_path / "real.docx"
        doc = docx_lib.Document()
        doc.add_paragraph("First sentence.")
        doc.add_paragraph("Second sentence!")
        doc.save(str(docx_file))

        result = read_text_from_file(str(docx_file))
        assert "First sentence." in result
        assert "Second sentence!" in result
