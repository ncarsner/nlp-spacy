import pytest
from unittest.mock import Mock, patch, MagicMock

import docx
import pypdf

from utils.functions import remove_punctuation
from utils.sentence_types import read_text_from_file


class TestFunctions:
    """Test suite for functions.py utilities"""

    def test_remove_punctuation_basic(self):
        """Test basic punctuation removal"""
        text = "Hello, world!"
        result = remove_punctuation(text)
        assert result == "Hello world"

    def test_remove_punctuation_all_types(self):
        """Test removal of all punctuation types"""
        text = "Hello! How are you? I'm fine, thanks. #winning @home"
        result = remove_punctuation(text)
        assert "!" not in result
        assert "?" not in result
        assert "," not in result
        assert "." not in result
        assert "'" not in result
        assert "#" not in result
        assert "@" not in result

    def test_remove_punctuation_empty_string(self):
        """Test with empty string"""
        result = remove_punctuation("")
        assert result == ""

    def test_remove_punctuation_no_punctuation(self):
        """Test with text that has no punctuation"""
        text = "Hello world"
        result = remove_punctuation(text)
        assert result == text

    def test_remove_punctuation_only_punctuation(self):
        """Test with only punctuation"""
        text = "!@#$%^&*()_+-={}[]|\\:;\"'<>,.?/"
        result = remove_punctuation(text)
        assert result == ""

    def test_remove_punctuation_preserves_numbers(self):
        """Test that numbers are preserved"""
        text = "Price: $100.50"
        result = remove_punctuation(text)
        assert "100" in result
        assert "50" in result

    def test_remove_punctuation_preserves_whitespace(self):
        """Test that whitespace is preserved"""
        text = "Hello,   world!  How are you?"
        result = remove_punctuation(text)
        assert "   " in result  # Multiple spaces preserved
        assert "  " in result


class TestReadTextFromFile:
    """Test suite for read_text_from_file() in sentence_types.py"""

    def test_read_txt_file(self, tmp_path):
        """Test reading a plain text file with UTF-8 encoding"""
        txt_file = tmp_path / "sample.txt"
        txt_file.write_text("Hello world. How are you?", encoding='utf-8')
        result = read_text_from_file(str(txt_file))
        assert result == "Hello world. How are you?"

    def test_read_txt_file_utf8_content(self, tmp_path):
        """Test reading a plain text file containing non-ASCII UTF-8 characters"""
        txt_file = tmp_path / "utf8.txt"
        txt_file.write_text("Héllo wörld!", encoding='utf-8')
        result = read_text_from_file(str(txt_file))
        assert result == "Héllo wörld!"

    def test_file_not_found_exits(self, tmp_path):
        """Test that a missing file causes SystemExit"""
        with pytest.raises(SystemExit):
            read_text_from_file(str(tmp_path / "nonexistent.txt"))

    @patch('utils.sentence_types.Path')
    def test_read_docx_file(self, mock_path_cls):
        """Test reading a .docx file extracts paragraph text"""
        mock_path = MagicMock()
        mock_path.exists.return_value = True
        mock_path.suffix = '.docx'
        mock_path.__str__ = lambda self: 'document.docx'
        mock_path_cls.return_value = mock_path

        mock_para1 = Mock()
        mock_para1.text = "Hello world."
        mock_para2 = Mock()
        mock_para2.text = "How are you?"

        mock_doc = Mock()
        mock_doc.paragraphs = [mock_para1, mock_para2]

        with patch('utils.sentence_types.docx.Document', return_value=mock_doc):
            result = read_text_from_file('document.docx')

        assert result == "Hello world.\nHow are you?"

    @patch('utils.sentence_types.Path')
    def test_read_docx_error_exits(self, mock_path_cls):
        """Test that a docx reading error causes SystemExit"""
        mock_path = MagicMock()
        mock_path.exists.return_value = True
        mock_path.suffix = '.docx'
        mock_path.__str__ = lambda self: 'bad.docx'
        mock_path_cls.return_value = mock_path

        with patch('utils.sentence_types.docx.Document', side_effect=Exception("corrupt file")):
            with pytest.raises(SystemExit):
                read_text_from_file('bad.docx')

    def test_read_docx_real_file(self, tmp_path):
        """Test reading an actual .docx file created with python-docx"""
        docx_file = tmp_path / "real.docx"
        doc = docx.Document()
        doc.add_paragraph("First sentence.")
        doc.add_paragraph("Second sentence!")
        doc.save(str(docx_file))

        result = read_text_from_file(str(docx_file))
        assert "First sentence." in result
        assert "Second sentence!" in result

    @patch('utils.sentence_types.Path')
    def test_read_pdf_file(self, mock_path_cls):
        """Test reading a .pdf file extracts text from all pages"""
        mock_path = MagicMock()
        mock_path.exists.return_value = True
        mock_path.suffix = '.pdf'
        mock_path.__str__ = lambda self: 'document.pdf'
        mock_path_cls.return_value = mock_path

        mock_page1 = Mock()
        mock_page1.extract_text.return_value = "Hello world."
        mock_page2 = Mock()
        mock_page2.extract_text.return_value = "How are you?"

        mock_reader = Mock()
        mock_reader.pages = [mock_page1, mock_page2]

        with patch('utils.sentence_types.pypdf.PdfReader', return_value=mock_reader):
            result = read_text_from_file('document.pdf')

        assert result == "Hello world.\nHow are you?"

    @patch('utils.sentence_types.Path')
    def test_read_pdf_skips_empty_pages(self, mock_path_cls):
        """Test that pages returning None or empty text are skipped"""
        mock_path = MagicMock()
        mock_path.exists.return_value = True
        mock_path.suffix = '.pdf'
        mock_path.__str__ = lambda self: 'document.pdf'
        mock_path_cls.return_value = mock_path

        mock_page1 = Mock()
        mock_page1.extract_text.return_value = "Hello world."
        mock_page2 = Mock()
        mock_page2.extract_text.return_value = None
        mock_page3 = Mock()
        mock_page3.extract_text.return_value = "Final sentence."

        mock_reader = Mock()
        mock_reader.pages = [mock_page1, mock_page2, mock_page3]

        with patch('utils.sentence_types.pypdf.PdfReader', return_value=mock_reader):
            result = read_text_from_file('document.pdf')

        assert result == "Hello world.\nFinal sentence."

    @patch('utils.sentence_types.Path')
    def test_read_pdf_error_exits(self, mock_path_cls):
        """Test that a PDF reading error causes SystemExit"""
        mock_path = MagicMock()
        mock_path.exists.return_value = True
        mock_path.suffix = '.pdf'
        mock_path.__str__ = lambda self: 'bad.pdf'
        mock_path_cls.return_value = mock_path

        with patch('utils.sentence_types.pypdf.PdfReader', side_effect=Exception("encrypted or corrupted")):
            with pytest.raises(SystemExit):
                read_text_from_file('bad.pdf')
